"""
将知乎图文最终版 .md 转换为可导入知乎编辑器的 .docx 文件。

功能：
1. 移除 YAML frontmatter
2. 将 Obsidian ![[文件名]] 语法转为标准 Markdown ![](截图/文件名)
3. 通过 Pandoc 转为 docx
4. 后处理：统一所有文本字体为微软雅黑，英文用 Arial

用法：
    python md_to_zhihu_docx.py <最终版.md路径>

输出：同目录下生成 <标题>-知乎发布版.docx

依赖：pandoc (系统), python-docx (pip)
"""

import re
import subprocess
import sys
import tempfile
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn


FONT_CN = "微软雅黑"
FONT_EN = "Arial"
FONT_CODE = "Consolas"

CODE_STYLE_NAMES = {
    "Source Code",
    "Verbatim Char",
    "Source Code Char",
}


def strip_frontmatter(content: str) -> str:
    """移除 YAML frontmatter（--- ... --- 块）"""
    if content.startswith("---"):
        end = content.find("---", 3)
        if end != -1:
            content = content[end + 3:].lstrip("\n")
    return content


def convert_wikilink_images(content: str) -> str:
    """将 ![[文件名.jpg]] 转为 ![](截图/文件名.jpg)"""
    return re.sub(
        r"!\[\[([^\]]+\.(jpg|png|jpeg|gif|webp))\]\]",
        r"![](截图/\1)",
        content,
    )


ARABIC_TO_CN = {
    "1": "一",
    "2": "二",
    "3": "三",
    "4": "四",
    "5": "五",
    "6": "六",
    "7": "七",
    "8": "八",
    "9": "九",
    "10": "十",
}


def fix_heading_numbers(content: str) -> str:
    """将标题中的阿拉伯数字编号转为中文编号，避免被知乎解析为有序列表。
    例如：## 1. Codex → ## 一、Codex
    """
    def _replace(m):
        prefix = m.group(1)  # ## / ### 等
        num = m.group(2)     # 阿拉伯数字
        rest = m.group(3)    # 标题正文
        cn = ARABIC_TO_CN.get(num, num)
        return f"{prefix}{cn}、{rest}"

    return re.sub(
        r"^(#{1,6}\s+)(\d{1,2})\.\s+(.+)$",
        _replace,
        content,
        flags=re.MULTILINE,
    )


def pandoc_md_to_docx(md_path: Path, docx_path: Path) -> None:
    """调用 Pandoc 将 Markdown 转为 docx"""
    cmd = [
        "pandoc",
        str(md_path),
        "-o", str(docx_path),
        "--from", "markdown",
        f"--resource-path={md_path.parent}",
    ]
    subprocess.run(cmd, check=True)


def is_code_style(run) -> bool:
    """判断 run 是否属于代码样式"""
    style = run.style
    if style and style.name in CODE_STYLE_NAMES:
        return True
    return False


def unify_fonts(docx_path: Path) -> None:
    """后处理：统一 docx 中所有文本的字体"""
    doc = Document(str(docx_path))

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            font = run.font
            if is_code_style(run):
                font.name = FONT_CODE
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CODE)
            else:
                font.name = FONT_EN
                run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = FONT_EN
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    doc.save(str(docx_path))


def main():
    if len(sys.argv) < 2:
        print("用法: python md_to_zhihu_docx.py <最终版.md路径>")
        sys.exit(1)

    src = Path(sys.argv[1])
    if not src.exists():
        print(f"文件不存在: {src}")
        sys.exit(1)

    content = src.read_text(encoding="utf-8")
    content = strip_frontmatter(content)
    content = convert_wikilink_images(content)
    content = fix_heading_numbers(content)

    stem = src.stem.replace("-知乎图文最终版", "")
    docx_name = f"{stem}-知乎发布版.docx"
    docx_path = src.parent / docx_name

    with tempfile.NamedTemporaryFile(
        mode="w", suffix=".md", encoding="utf-8", delete=False, dir=src.parent
    ) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)

    try:
        print(f"Pandoc 转换中: {tmp_path.name} -> {docx_name}")
        pandoc_md_to_docx(tmp_path, docx_path)
        print("统一字体中...")
        unify_fonts(docx_path)
        print(f"完成: {docx_path}")
    finally:
        tmp_path.unlink(missing_ok=True)


if __name__ == "__main__":
    main()
